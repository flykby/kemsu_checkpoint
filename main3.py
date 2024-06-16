import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io

# Устанавливаем seed для воспроизводимости результатов
np.random.seed(0)

# Создаем 12 случайных значений узлов сетки в диапазоне от 0 до 10
xn = np.sort(np.random.uniform(0, 10, 12))

# Создаем случайные значения функции в узлах сетки с плавающей точкой до двух знаков
yn = np.round(np.random.uniform(-10, 10, 12), 2)

# Аппроксимация методом наименьших квадратов (МНК)
degrees = [1, 2, 3, 4, 5]

# Функция для метода Гаусса
def gaussian_elimination(A, b):
    n = len(b)
    # Прямой ход
    for i in range(n):
        max_row = max(range(i, n), key=lambda r: abs(A[r][i]))
        if A[max_row][i] == 0:
            raise ValueError("Матрица вырождена")
        
        A[[i, max_row]] = A[[max_row, i]]
        b[[i, max_row]] = b[[max_row, i]]
        
        for j in range(i + 1, n):
            factor = A[j][i] / A[i][i]
            A[j, i:] -= factor * A[i, i:]
            b[j] -= factor * b[i]

    x = np.zeros_like(b)
    for i in range(n - 1, -1, -1):
        x[i] = (b[i] - np.dot(A[i, i+1:], x[i+1:])) / A[i, i]
    return x

# Определение матрицы Вандермонда и применение метода Гаусса для каждого полинома
x_vals = np.linspace(min(xn), max(xn), 1000)

plt.figure(figsize=(12, 8))
plt.scatter(xn, yn, color='red', label='Узлы сетки')

for d in degrees:
    coefficients_mnk = np.polyfit(xn, yn, d)
    poly_mnk = np.poly1d(coefficients_mnk)
    y_mnk = poly_mnk(x_vals)
    
    plt.plot(x_vals, y_mnk, label=f'МНК (степень {d})')

plt.legend()
plt.xlabel('x')
plt.ylabel('y')
plt.title('Аппроксимирующие функции методом наименьших квадратов')
plt.grid(True)
plt.show()

# Сохранение графиков в файл .docx
doc = Document()
doc.add_heading('Аппроксимирующие функции', 0)

for d in degrees:
    coefficients_mnk = np.polyfit(xn, yn, d)
    poly_mnk = np.poly1d(coefficients_mnk)
    y_mnk = poly_mnk(x_vals)
    
    plt.figure(figsize=(12, 8))
    plt.scatter(xn, yn, color='red', label='Узлы сетки')
    plt.plot(x_vals, y_mnk, label=f'МНК (степень {d})')
    plt.legend()
    plt.xlabel('x')
    plt.ylabel('y')
    plt.title(f'Аппроксимирующая функция МНК (степень {d})')
    plt.grid(True)
    
    buf = io.BytesIO()
    plt.savefig(buf, format='png')
    buf.seek(0)
    doc.add_picture(buf, width=Inches(6))
    buf.close()

doc.save('interpolation_approximation.docx')
